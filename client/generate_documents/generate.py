from docx import Document
import re


def replace_text(paragraph, key, value):
    if key in paragraph.text:
        paragraph.text = paragraph.text.replace(key, value)


def generate(data_arg: dict,  components_price: float) -> dict:
    data = {
        "{{NUMBER}}": data_arg["data"][0]["id"],
        "{{DATE}}": data_arg["data"][0]["created_at"].split(".")[0].replace("T", " ")[:-3],
        "{{FROM}}": data_arg["data"][0]["user"]["fio"],
        "{{TO}}": data_arg["data"][0]["client"]["fio"],
        "{{GLOBAL_COUNT}}": len(data_arg["data"][0]["components_replied"]),
        "{{GLOBAL_PRICE}}": components_price,
        "components": data_arg["data"][0]["components_replied"]
    }
    pattern = r"{{\w{1,}}}"
    template = 'generate_documents/templates/template.docx'
    result = f'накладная_{data["{{NUMBER}}"]}.docx'.replace(" ", "_")
    template_doc = Document(template)
    try:
        for k, v in data.items():
            if re.fullmatch(pattern, k.strip()):
                for paragraph in template_doc.paragraphs:
                    replace_text(paragraph, k, str(v))
        my_table = template_doc.tables[0]
        n = 0
        for line in data["components"]:
            i = 0
            row_data = [n + 1, line["name"], line["type"], line["price"]]
            n += 1
            row = my_table.add_row()
            for cell in row.cells:
                cell.text = str(row_data[i])
                i += 1
        template_doc.save(str(f"documents/{result}"))

    except Exception as ex:
        print(ex)
        return {"status": "error", "message": ex}
    return {"status": "ok", "message": "Document created!", "data": result}
